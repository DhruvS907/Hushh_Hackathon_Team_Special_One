import os
import sys
import json
import re
from datetime import datetime
from typing import Dict, List, Optional, Tuple, Annotated, Sequence, Any
from enum import Enum
from dataclasses import dataclass
from langchain_openai import ChatOpenAI
from langchain_core.messages import BaseMessage, HumanMessage, AIMessage
from langgraph.graph import StateGraph, START, END
from langgraph.graph.message import add_messages
from typing import TypedDict
from dotenv import load_dotenv

# --- Corrected Path Setup ---
project_root = os.path.abspath(os.path.join(os.path.dirname(__file__), "..", "..", ".."))
sys.path.append(project_root)
backend_root = os.path.abspath(os.path.join(os.path.dirname(__file__), ".."))
sys.path.append(backend_root)

# --- RAG Imports ---
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_community.vectorstores import FAISS
from langchain_google_genai import GoogleGenerativeAIEmbeddings
from langchain_core.documents import Document
from langchain_core.vectorstores import VectorStoreRetriever

# --- Local & Library Imports ---
from agents.info_responder_agent import info_responder_agent
from agents.schedular_agent import calendar_agent
from Email_Summarizer import fetch_user_sent_emails
from hushh_mcp.consent.token import validate_token
from hushh_mcp.constants import ConsentScope

# --- NEW: Imports for PDF and DOCX processing ---
import pypdf
import docx

load_dotenv()

# --- Enums and Dataclasses ---
class AgentType(Enum):
    SCHEDULER = "scheduler"
    INFO_RESPONDER = "info_responder"
    GENERAL_RESPONDER = "general_responder"
    NO_RESPONSE = "no_response"

@dataclass
class EmailContext:
    subject: str
    sender: str
    sender_email: str
    body: str
    summary: str
    intent: str
    snippet: str

@dataclass
class ResponsePlan:
    agent_type: AgentType
    confidence: float
    reasoning: str
    suggested_action: str
    requires_user_input: bool = False

# --- LangGraph State ---
class EmailState(TypedDict):
    """State for the email orchestration workflow"""
    messages: Annotated[Sequence[BaseMessage], add_messages]
    email_context: EmailContext
    user_email: str
    user_name: str
    user_suggestion: Optional[str]
    document_content: Optional[bytes]
    document_filename: Optional[str]
    history_retriever: Optional[VectorStoreRetriever]
    tone_retriever: Optional[VectorStoreRetriever]
    knowledge_retriever: Optional[VectorStoreRetriever]
    has_kb_consent: bool
    attachment_to_send: Optional[Dict[str, Any]]
    response_plan: Optional[ResponsePlan]
    agent_outcome: Optional[str]
    final_response: Optional[str]
    error: Optional[str]
    conversation_history: Optional[List[str]]

# --- Main Orchestration Agent ---
class OrchestrationAgent:
    def __init__(self, user_name: str, user_email: str, access_token: str):
        self.user_email = user_email
        self.user_name = user_name
        self.access_token = access_token
        self.llm = ChatOpenAI(
            openai_api_key=os.environ["GROQ_API_KEY"],
            openai_api_base="https://api.groq.com/openai/v1",
            model="qwen/qwen3-32b",
            temperature=0.3,
        )
        # ✅ Single embeddings model for all retrieval tasks
        self.embeddings = GoogleGenerativeAIEmbeddings(
            model="models/embedding-001",
            google_api_key=os.environ.get("GOOGLE_API_KEY")
        )
        self.intent_mapping = {
            "Scheduling or rescheduling a meeting or event": AgentType.SCHEDULER,
            "Requesting information or clarification": AgentType.INFO_RESPONDER,
            "Marketing emails or newsletters": AgentType.NO_RESPONSE,
            "Informational only – no action required (FYI)": AgentType.NO_RESPONSE,
            "Announcing a new product or feature": AgentType.NO_RESPONSE,
            "Shipping, delivery, or order tracking update": AgentType.NO_RESPONSE,
        }
        self.workflow = self._build_workflow()

    # --- MODIFIED: This function now supports PDF and DOCX files ---
    def _build_knowledge_retriever(self, user_email: str) -> Optional[VectorStoreRetriever]:
        """Scans a user-specific directory for .pdf, .docx, .txt, and .md files and builds a searchable retriever."""
        if not user_email:
            return None

        base_path = os.path.join(os.path.dirname(__file__), "..", "user_knowledge_bases")
        sanitized_email = user_email.replace("@", "_at_").replace(".", "_dot_")
        user_knowledge_base_path = os.path.join(base_path, sanitized_email)

        if not os.path.exists(user_knowledge_base_path):
            return None

        docs = []
        for filename in os.listdir(user_knowledge_base_path):
            file_path = os.path.join(user_knowledge_base_path, filename)
            content = ""
            if not os.path.isfile(file_path):
                continue
            
            try:
                # Process PDF files
                if filename.lower().endswith('.pdf'):
                    with open(file_path, 'rb') as f:
                        reader = pypdf.PdfReader(f)
                        for page in reader.pages:
                            content += page.extract_text() or ""
                
                # Process DOCX files
                elif filename.lower().endswith('.docx'):
                    doc = docx.Document(file_path)
                    for para in doc.paragraphs:
                        content += para.text + "\n"

                # Process TXT and MD files
                elif filename.lower().endswith(('.txt', '.md')):
                    with open(file_path, 'r', encoding='utf-8') as f:
                        content = f.read()
                
                # If content was extracted, create a Document
                if content:
                    docs.append(Document(page_content=content, metadata={"source": filename}))

            except Exception as e:
                print(f"Error processing file {filename} for user {user_email}: {e}")

        if not docs:
            return None

        text_splitter = RecursiveCharacterTextSplitter(chunk_size=1000, chunk_overlap=100)
        splits = text_splitter.split_documents(docs)
        vector_store = FAISS.from_documents(splits, self.embeddings)
        return vector_store.as_retriever(search_kwargs={"k": 3})

    def _build_workflow(self) -> StateGraph:
        workflow = StateGraph(EmailState)
        workflow.add_node("fetch_and_index_tone_emails", self._fetch_and_index_tone_emails_node)
        workflow.add_node("analyzer", self._analyze_email_node)
        workflow.add_node("scheduler_agent", self._scheduler_agent_node)
        workflow.add_node("info_agent", self._info_agent_node)
        workflow.add_node("general_agent", self._general_agent_node)
        workflow.add_node("no_response", self._no_response_node)
        workflow.add_node("composer", self._compose_final_email_node)

        workflow.add_edge(START, "fetch_and_index_tone_emails")
        workflow.add_edge("fetch_and_index_tone_emails", "analyzer")
        workflow.add_conditional_edges(
            "analyzer", self._route_to_agent,
            {
                "scheduler": "scheduler_agent", "info_responder": "info_agent",
                "general_responder": "general_agent", "no_response": "no_response"
            }
        )
        workflow.add_edge("scheduler_agent", "composer")
        workflow.add_edge("info_agent", "composer")
        workflow.add_edge("general_agent", "composer")
        workflow.add_edge("no_response", END)
        workflow.add_edge("composer", END)

        return workflow.compile()

    def _fetch_and_index_tone_emails_node(self, state: EmailState) -> EmailState:
        try:
            sent_emails = fetch_user_sent_emails(self.access_token, days=7)
            if not sent_emails:
                return {**state, "tone_retriever": None}
            documents = [Document(page_content=email['body'], metadata={'subject': email['subject']}) for email in sent_emails]
            text_splitter = RecursiveCharacterTextSplitter(chunk_size=1000, chunk_overlap=100)
            splits = text_splitter.split_documents(documents)
            vector_store = FAISS.from_documents(splits, self.embeddings)
            retriever = vector_store.as_retriever(search_kwargs={"k": 3})
            return {**state, "tone_retriever": retriever}
        except Exception as e:
            print(f"Error creating tone retriever: {e}")
            return {**state, "tone_retriever": None}

    def _analyze_email_node(self, state: EmailState) -> EmailState:
        email_context = state["email_context"]
        initial_agent = self.intent_mapping.get(email_context.intent, AgentType.GENERAL_RESPONDER)
        history_context = "\n".join(state.get("conversation_history", []))

        analysis_prompt = f"""
        Analyze this email and determine the best response strategy.
        Email Details:
        - Subject: {email_context.subject}
        - Sender: {email_context.sender}
        - Intent: {email_context.intent}
        - Summary: {email_context.summary}

        Conversation History:
        {history_context}

        Available Agents: SCHEDULER, INFO_RESPONDER, GENERAL_RESPONDER, NO_RESPONSE.
        Return only valid JSON:
        {{
            "agent_type": "scheduler|info_responder|general_responder|no_response",
            "confidence": 0.0-1.0,
            "reasoning": "Brief explanation of why this agent is chosen",
            "suggested_action": "What specific action should be taken"
        }}
        """
        try:
            response = self.llm.invoke([HumanMessage(content=analysis_prompt)])
            analysis = self._extract_json(response.content)
            response_plan = ResponsePlan(
                agent_type=AgentType(analysis["agent_type"]),
                confidence=analysis["confidence"],
                reasoning=analysis["reasoning"],
                suggested_action=analysis["suggested_action"]
            )
        except Exception as e:
            print(f"LLM analysis failed: {e}, falling back to intent mapping.")
            response_plan = ResponsePlan(
                agent_type=initial_agent,
                confidence=0.7,
                reasoning=f"Fallback based on intent: {email_context.intent}",
                suggested_action="Handle using the mapped agent"
            )

        analysis_msg = AIMessage(content=f"Analyzed email. Route to: {response_plan.agent_type.value}")
        return {**state, "messages": state["messages"] + [analysis_msg], "response_plan": response_plan}

    def _route_to_agent(self, state: EmailState) -> str:
        return state["response_plan"].agent_type.value if state["response_plan"] else "general_responder"

    def _scheduler_agent_node(self, state: EmailState) -> EmailState:
        email_context = state["email_context"]
        user_suggestion = state.get("user_suggestion")
        try:
            sender_email = self._extract_email_from_sender(email_context.sender)
            message_content = f"Email from {email_context.sender}:\n{email_context.body}\n{f'User suggestion: {user_suggestion}' if user_suggestion else ''}"
            initial_state = {'messages': [HumanMessage(content=message_content)], 'senders_email_address': sender_email, 'users_email_address': state["user_email"]}
            result = calendar_agent.invoke(initial_state)
            agent_outcome = result['messages'][-1].content
        except Exception as e:
            agent_outcome = f"Error in scheduler: {str(e)}"

        return {**state, "agent_outcome": agent_outcome}

    def _info_agent_node(self, state: EmailState) -> EmailState:
        email_context = state["email_context"]
        user_suggestion = state.get("user_suggestion")
        document_content = state.get("document_content")
        document_filename = state.get("document_filename")
        knowledge_retriever = state.get("knowledge_retriever")
        has_kb_consent = state.get("has_kb_consent", False)
        user_email = state["user_email"]

        attachment_to_send = None

        # If a file was uploaded for regeneration, automatically set it as the attachment.
        if document_content and document_filename:
            print(f"File '{document_filename}' provided for regeneration. Prioritizing as attachment.")
            attachment_to_send = {"filename": document_filename, "content": document_content}

        query = f"{email_context.summary}\n\n{f'User guidance: {user_suggestion}' if user_suggestion else ''}"

        active_retriever = None
        if has_kb_consent and knowledge_retriever:
            print(f"Knowledge base consent granted for {user_email}. Searching their files.")
            active_retriever = knowledge_retriever
        else:
            print(f"Knowledge base consent not granted for {user_email} or retriever not available. Skipping file search.")

        try:
            raw_outcome = info_responder_agent(
                query=query,
                doc_content=document_content,
                doc_filename=document_filename,
                knowledge_retriever=active_retriever
            )

            agent_outcome = raw_outcome
            attachment_match = re.search(r"\[ATTACH_FILE:\s*(.*?)\]", raw_outcome)

            # If the agent recommended a file from the KB and we haven't already set one from an upload
            if attachment_match and not attachment_to_send:
                attachment_filename = attachment_match.group(1).strip()
                agent_outcome = raw_outcome.replace(attachment_match.group(0), "").strip()

                base_path = os.path.join(os.path.dirname(__file__), "..", "user_knowledge_bases")
                sanitized_email = user_email.replace("@", "_at_").replace(".", "_dot_")
                user_knowledge_base_path = os.path.join(base_path, sanitized_email)
                file_path = os.path.join(user_knowledge_base_path, attachment_filename)

                if os.path.exists(file_path):
                    with open(file_path, 'rb') as f:
                        file_content = f.read()
                    attachment_to_send = {"filename": attachment_filename, "content": file_content}
                    print(f"Prepared '{attachment_filename}' for attachment for user {user_email}.")
                else:
                    print(f"Warning: Agent requested attachment '{attachment_filename}' for user {user_email}, but it was not found.")
                    agent_outcome += f"\n\n(Note: I was unable to find the requested attachment: {attachment_filename})"
            else:
                # Just strip the tag if it exists, as we're either ignoring it or there isn't one.
                agent_outcome = re.sub(r"\[ATTACH_FILE:\s*(.*?)\]", "", raw_outcome).strip()

            return {**state, "agent_outcome": agent_outcome, "attachment_to_send": attachment_to_send}
        except Exception as e:
            agent_outcome = f"Error processing info request: {str(e)}"
            return {**state, "agent_outcome": agent_outcome, "attachment_to_send": None}

    def _general_agent_node(self, state: EmailState) -> EmailState:
        email_context = state["email_context"]
        user_suggestion = state.get("user_suggestion")

        recipient_name = email_context.sender.split('<')[0].strip()
        if '@' in recipient_name:
            recipient_name = "there"

        history_context = "\n".join(state.get("conversation_history", []))

        response_prompt = f"""
        Respond to the following email professionally:

        From: {state['email_context'].sender}
        Subject: {state['email_context'].subject}
        Context: {state['email_context'].body[:500]}

        Conversation History:
        {history_context}

        {f"User guidance: {state['user_suggestion']}" if state.get("user_suggestion") else ""}
        """
        try:
            response = self.llm.invoke([HumanMessage(content=response_prompt)])
            agent_outcome = self._strip_think_block(response.content)
        except Exception as e:
            agent_outcome = f"Error generating response: {str(e)}"

        return {**state, "agent_outcome": agent_outcome}

    def _no_response_node(self, state: EmailState) -> EmailState:
        return {**state, "final_response": "This email doesn't require a response."}

    def _compose_final_email_node(self, state: EmailState) -> EmailState:
        agent_outcome = state.get("agent_outcome", "No information was generated.")
        email_context = state["email_context"]
        tone_retriever = state.get("tone_retriever")

        recipient_name = email_context.sender.split('<')[0].strip()
        if '@' in recipient_name:
            recipient_name = "there"

        tone_examples = ""
        if tone_retriever:
            try:
                similar_docs = tone_retriever.invoke(email_context.body)
                if similar_docs:
                    tone_examples += "Please use a similar tone and style to the following examples from past emails sent by the user:\n\n"
                    for i, doc in enumerate(similar_docs):
                        tone_examples += f"--- Example {i+1} ---\n{doc.page_content}\n\n"
            except Exception as e:
                print(f"Could not retrieve tone examples: {e}")

        response_prompt = f"""
        You are an AI assistant writing a professional email on behalf of {self.user_name}.

        Your task is to compose a final email response.
        Use the "Contextual Information" below to formulate your answer.
        Integrate this information naturally into a polite and helpful email.

        {tone_examples}

        **Contextual Information to Use for the Reply:**
        ---
        {agent_outcome}
        ---

        **Original Email Details:**
        - From: {email_context.sender}
        - Subject: {email_context.subject}

        **Instructions:**
        1. Address the email to '{recipient_name}'.
        2. Write a complete and professional email response.
        3. Sign the email with the name '{self.user_name}'.

        GIVE YOUR RESPONSE ONLY THE BODY OF THE EMAIL AND NOTHING ELSE AND DONT GIVE ANY WORDS IN BOLD
        """
        response = self.llm.invoke([HumanMessage(content=response_prompt)])
        final_response = self._strip_think_block(response.content)

        return {**state, "final_response": final_response}

    def _strip_think_block(self, text: str) -> str:
        return re.sub(r"^<think>.*?</think>", "", text, flags=re.DOTALL).strip()

    def generate_response(self, email_context: EmailContext, consent_token: str, user_suggestion: Optional[str] = None, document_content: Optional[bytes] = None, document_filename: Optional[str] = None, conversation_history: Optional[List[str]] = None, knowledge_base_consent_token: Optional[str] = None) -> Dict:
        is_valid, reason, parsed_token = validate_token(consent_token, expected_scope=ConsentScope.VAULT_READ_EMAIL)
        if not is_valid:
            raise PermissionError(f"Consent validation failed: {reason}")
        if parsed_token.user_id != self.user_email:
            raise PermissionError("User ID in token does not match")

        has_kb_consent = False
        if knowledge_base_consent_token:
            is_kb_valid, _, kb_parsed_token = validate_token(knowledge_base_consent_token, expected_scope=ConsentScope.KNOWLEDGE_BASE_READ)
            if is_kb_valid and kb_parsed_token.user_id == self.user_email:
                has_kb_consent = True

        knowledge_retriever = self._build_knowledge_retriever(self.user_email)

        initial_state = {
            "messages": [HumanMessage(content=f"Processing email: {email_context.subject}")],
            "email_context": email_context,
            "user_email": self.user_email,
            "user_name": self.user_name,
            "user_suggestion": user_suggestion,
            "document_content": document_content,
            "document_filename": document_filename,
            "conversation_history": conversation_history,
            "knowledge_retriever": knowledge_retriever,
            "has_kb_consent": has_kb_consent,
            "attachment_to_send": None,
        }
        try:
            final_state = self.workflow.invoke(initial_state)
            response_plan = final_state.get("response_plan")
            final_response = final_state.get("final_response", "No response generated")
            attachment = final_state.get("attachment_to_send")

            agent_type = response_plan.agent_type.value if response_plan else "general_responder"
            if agent_type == "no_response":
                return {"response_type": "no_response", "message": final_response, "reasoning": "N/A", "confidence": 1.0, "attachment": None}

            return {
                "response_type": agent_type, "message": final_response,
                "reasoning": response_plan.reasoning if response_plan else "Default processing",
                "confidence": response_plan.confidence if response_plan else 0.7,
                "attachment": attachment,
            }
        except Exception as e:
            return {"response_type": "error", "message": f"Error: {str(e)}", "reasoning": "System error", "confidence": 0.0, "attachment": None}

    def _extract_email_from_sender(self, sender: str) -> str:
        email_match = re.search(r'<([^>]+)>', sender)
        return email_match.group(1) if email_match else sender.strip()

    def _extract_json(self, response_text: str) -> Optional[Dict]:
        match = re.search(r"\{.*\}", response_text, re.DOTALL)
        if match:
            try:
                return json.loads(match.group())
            except json.JSONDecodeError:
                return None
        return None


def process_email_with_orchestration(email_data: Dict, user_email: str, user_name: str, consent_token: str, access_token: str, user_suggestion: Optional[str] = None, document_content: Optional[bytes] = None, document_filename: Optional[str] = None, conversation_history: Optional[List[str]] = None, knowledge_base_consent_token: Optional[str] = None) -> Dict:
    email_context = EmailContext(
        subject=email_data.get('subject', ''),
        sender=email_data.get('sender', ''),
        sender_email=email_data.get('sender', ''),
        body=email_data.get('body', ''),
        summary=email_data.get('summary', ''),
        intent=email_data.get('intent', ''),
        snippet=email_data.get('snippet', '')
    )
    orchestrator = OrchestrationAgent(user_name, user_email, access_token)
    return orchestrator.generate_response(
        email_context, consent_token, user_suggestion, document_content,
        document_filename, conversation_history, knowledge_base_consent_token
    )